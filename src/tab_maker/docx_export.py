"""DOCX export for Tab-Maker songs."""
from __future__ import annotations

from pathlib import Path
from typing import Union

from .models import Song
from .text import song_to_plain_lines

try:  # pragma: no cover - optional dependency
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:  # pragma: no cover
    Document = None  # type: ignore
    Pt = None  # type: ignore
    WD_ALIGN_PARAGRAPH = None  # type: ignore


def _ensure_docx_available() -> None:
    if Document is None:
        raise RuntimeError(
            "python-docx is required for DOCX export. Install with 'pip install python-docx'."
        )


def song_to_docx(song: Song, destination: Union[str, Path]) -> Path:
    """Write the song to a DOCX file and return the output path."""
    _ensure_docx_available()
    assert Document is not None

    document = Document()
    metadata = song.metadata

    title = metadata.get("title")
    artist = metadata.get("artist")
    if title and artist:
        header_text = f"{title} - {artist}"
    else:
        header_text = title or artist

    if header_text:
        section = document.sections[0]
        header = section.header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.text = ""
        try:
            paragraph.style = document.styles["Title"]
        except KeyError:
            pass
        run = paragraph.add_run(header_text)
        run.bold = True
        run.font.name = "Courier New"
        if Pt is not None:
            run.font.size = Pt(16)
        if WD_ALIGN_PARAGRAPH is not None:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if "title" in metadata:
        document.core_properties.title = metadata["title"]
    if "artist" in metadata:
        document.core_properties.author = metadata["artist"]

    for line in song_to_plain_lines(song):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        if Pt is not None:
            run.font.size = Pt(11)

    output_path = Path(destination)
    document.save(output_path)
    return output_path


__all__ = ["song_to_docx"]
